from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# Function to extract data from the file
def extract_data(file_path):
    """
    Extracts data from a text file and returns a list of tuples containing the data.
    :param file_path: The path to the file to extract data from.
    :return: A list of tuples containing the extracted data based on the patters in the file.
    """
    with open(file_path, 'r') as file:
        content = file.read()

    # Use regular expressions to find the data patterns
    name_pattern = r'Name: (.+)'
    age_pattern = r'Age: (\d+)'
    email_pattern = r'Email: (.+@.+\..+)'

    names = re.findall(name_pattern, content)
    ages = re.findall(age_pattern, content)
    emails = re.findall(email_pattern, content)

    data = list(zip(names, ages, emails))
    return data


# Function to set table borders
def set_table_borders(table):
    """
    Sets the borders of a table.
    :param table: The table to set the borders for.
    :return: None
    """
    # get the underlying XML element of the table
    tbl = table._element
    # create a new element for the table borders
    tblBorders = OxmlElement('w:tblBorders')
    # set the attributes for the table
    # XML element is created and configured with attributes for value (single), size (4), space (0), and color (000000).
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        # set the border style for a table border
        border.set(qn('w:val'), 'single') # 'single' specifies that the border style should be a single line.
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)


# Function to write data to a Word document
def write_to_word(data, output_path):
    doc = Document()
    doc.add_heading('Extracted Data Table', level=1)

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'Email'

    for name, age, email in data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = age
        row_cells[2].text = email

    # Add borders to the table
    set_table_borders(table)

    doc.save(output_path)


# Main program
if __name__ == '__main__':
    input_file = 'sample_data.txt'
    output_file = 'extracted_data_with_borders.docx'

    extracted_data = extract_data(input_file)
    write_to_word(extracted_data, output_file)

    print(f'Data extracted and written to {output_file}')
